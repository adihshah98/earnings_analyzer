"""Extract plain text from Word (.docx) documents for transcript ingestion."""

import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a .docx file's binary content.

    Uses python-docx to read paragraphs and tables in document order, joined
    by newlines. Suitable for earnings call transcripts.

    Args:
        file_content: Raw bytes of the .docx file.

    Returns:
        Extracted text, with paragraphs separated by newlines.

    Raises:
        ValueError: If the content is not valid .docx or contains no text.
    """
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            "python-docx is required for Word document upload. "
            "Install with: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(file_content))
    except Exception as e:
        raise ValueError(f"Invalid or corrupted Word document: {e}") from e

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    if not parts:
        raise ValueError("Document contains no extractable text.")

    return "\n\n".join(parts)
